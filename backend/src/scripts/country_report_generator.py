"""
Web Archiver & Scraper (Multithreaded)
Pipeline:
  Step 1 — Collect URLs (Bloomberg sitemap + Gemini web search for other sources)
  Step 2 — Scrape each URL via archive.today → Gemini url_context → BeautifulSoup fallback
  Step 3 — report_generator(): synthesise all content into a structured report via Gemini
  Step 4 — Save report as dated .docx in /reports
"""

from concurrent.futures._base import Future
import json
import time
import re
import os

import random
import threading
from typing import Any
import requests
import subprocess
import xml.etree.ElementTree as ET
from dataclasses import dataclass
from datetime import datetime, timezone
from pathlib import Path
from urllib.parse import urlparse
from concurrent.futures import ThreadPoolExecutor, as_completed
from threading import Lock
from bs4 import BeautifulSoup
from dotenv import load_dotenv
from google import genai
from google.genai import types
from docx import Document as DocxDocument
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import font_hanken_grotesk

load_dotenv()

_COUNTRY = "Brazil"
_MONTH_YEAR = datetime.now(timezone.utc).strftime("%B %Y")

# Font names Word recognises
_FONT_REGULAR = "Hanken Grotesk"
_FONT_MEDIUM = "Hanken Grotesk Medium"
_FONT_BOLD = "Hanken Grotesk Bold"


USER_AGENTS = [
    "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/122.0.0.0 Safari/537.36",
    "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/121.0.0.0 Safari/537.36",
    "Mozilla/5.0 (X11; Linux x86_64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36",
    "Mozilla/5.0 (Windows NT 10.0; Win64; x64; rv:123.0) Gecko/20100101 Firefox/123.0",
    "Mozilla/5.0 (Macintosh; Intel Mac OS X 14_3) AppleWebKit/605.1.15 (KHTML, like Gecko) Version/17.2 Safari/605.1.15",
]

# ─────────────────────────────────────────────
#  CONFIG
# ─────────────────────────────────────────────
MAX_WORKERS = 3
REQUESTS_PER_SECOND = 0.4
GEMINI_MODEL = "gemini-2.5-flash"
MAX_RETRIES = 5
RETRY_DELAY = 2
REPORTS_DIR = Path("reports")

# Sources for Gemini search (non-Bloomberg)
SEARCH_SOURCES = [
    ####### Global #######
    "BBC",
    "Al Jazeera",
    "Financial Times",
    ####### China #######
    # "South China Morning Post",
    # "Reuters",
    # "BBC China",
    # "Global Times",
    # "People's Daily",
    # "Xinhua",
    # "CGTN",
    ####### Japan #######
    # "Politico",
    # "NHK World",  # Japan's public broadcaster (English)
    # "The Japan Times",
    # "Mainichi Shimbun",
    # "Yomiuri Shimbun",
    ####### Turkey #######
    # "Hürriyet Daily News",
    # "Daily Sabah",
    # "Turkish Minute",
    # "BBC Türkçe:",
    ####### Australia #######
    # "ABC",
    # "SBS",
    # "The Australian",
    ####### South Korea #######
    # "Yonhap News Agency",
    # "Korea Herald",
    # "Korea JoongAng Daily",
    # "The Korea Times",
    ####### Italy #######
    # "Folha de S.Paulo",
    # "O Globo",
    # "O Estado de S. Paulo ",
    # "Valor Econômico ",
    # "BBC Brasil",
    ####### South Africa #######
    # "Daily Maverick",
    # "Business Day",
    # "Mail & Guardian",
    # "Sunday Times",
    # "TimesLive",
    ####### Indonesia #######
    # "Antara",
    # "TVRI",
    # "Metro TV ",
    # "Kompas TV",
    # "CNN Indonesia",
    # "BBC Indonesia",
    # "Reuters",
    ####### Brazil #######
    "Folha de S.Paulo",
    "O Globo",
    "O Estado de S. Paulo ",
    "Valor Econômico ",
    "BBC Brasil",
]

BLOCKED_DOMAINS = {
    "instagram",
    "facebook",
    "twitter",
    "tiktok",
    "x.com",
    "youtu.be",
    "youtube",
}

# ─────────────────────────────────────────────
#  Thread-safe print
# ─────────────────────────────────────────────
_lock = Lock()


def tprint(*args, **kwargs):
    with _lock:
        print(*args, **kwargs)


# ─────────────────────────────────────────────
#  Global Rate Limiter + Backoff
# ─────────────────────────────────────────────
class RateLimiter:
    def __init__(self, calls_per_second: float):
        self.interval = 1.0 / calls_per_second
        self._lock = threading.Lock()
        self._last_call = 0.0

    def wait(self):
        with self._lock:
            now = time.monotonic()
            wait = self.interval - (now - self._last_call)
            if wait > 0:
                time.sleep(wait)
            self._last_call = time.monotonic()


class GlobalBackoff:
    def __init__(self):
        self._lock = threading.Lock()
        self._backoff_until = 0.0

    def trigger(self, wait_seconds: float):
        with self._lock:
            self._backoff_until = time.monotonic() + wait_seconds
        tprint(f"  [429] Global backoff — all threads pausing {wait_seconds:.1f}s")

    def wait_if_needed(self):
        while True:
            with self._lock:
                remaining = self._backoff_until - time.monotonic()
            if remaining <= 0:
                break
            time.sleep(min(remaining, 0.5))


_rate_limiter = RateLimiter(calls_per_second=REQUESTS_PER_SECOND)
_global_backoff = GlobalBackoff()


# ─────────────────────────────────────────────
#  Gemini client + tool declarations
# ─────────────────────────────────────────────
_client = genai.Client(api_key=os.environ.get("GEMINI_API_KEY"))

_BEAUTIFULSOUP_TOOL = types.Tool(
    function_declarations=[
        types.FunctionDeclaration(
            name="scrape_with_beautifulsoup",
            description=(
                """Fetches a URL and extracts the main article text using BeautifulSoup. 
                Call this when url_context cannot retrieve content from the page (e.g. paywalls, bot-detection, or empty responses). 
                Return only the real article text — do not fabricate content."""
            ),
            parameters=types.Schema(
                type=types.Type.OBJECT,
                properties={
                    "url": types.Schema(
                        type=types.Type.STRING,
                        description="The full URL to fetch and scrape.",
                    )
                },
                required=["url"],
            ),
        )
    ]
)

# ── Scraping configs — MUST be separate requests.
# Gemini forbids mixing url_context with FunctionDeclarations in one request.
#
#  Pass 1: url_context only  — Gemini reads the page directly
#  Pass 2: BS4 function tool — used if Pass 1 returns empty content
_scrape_url_context_config = types.GenerateContentConfig(tools=[{"url_context": {}}])
_scrape_bs4_config = types.GenerateContentConfig(tools=[_BEAUTIFULSOUP_TOOL])

# Config for URL collection (web search enabled)
_search_config = types.GenerateContentConfig(
    tools=[{"google_search": {}}, {"url_context": {}}],
)

# Config for report writing (no tools needed — pure synthesis)
_report_config = types.GenerateContentConfig()


# ─────────────────────────────────────────────
#  Data model
# ─────────────────────────────────────────────
@dataclass
class SitemapEntry:
    url: str
    last_modified: datetime
    change_freq: str
    priority: float
    section: str
    date: str
    slug: str


# ─────────────────────────────────────────────
#  STEP 1A — Bloomberg sitemap parser (unchanged logic)
# ─────────────────────────────────────────────
def parse_bloomberg_sitemap(xml_url: str) -> list[SitemapEntry]:
    headers = {
        "User-Agent": "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) "
        "AppleWebKit/537.36 (KHTML, like Gecko) Chrome/122.0.0.0 Safari/537.36",
        "Referer": "https://www.bloomberg.com/",
    }
    response = requests.get(xml_url, headers=headers, timeout=20)
    response.raise_for_status()

    namespace = {"sm": "http://www.sitemaps.org/schemas/sitemap/0.9"}
    root = ET.fromstring(response.content)
    entries = []

    for url_el in root.findall("sm:url", namespace):
        loc = url_el.findtext("sm:loc", default="", namespaces=namespace)
        lastmod = url_el.findtext("sm:lastmod", default="", namespaces=namespace)
        changefreq = url_el.findtext("sm:changefreq", default="", namespaces=namespace)
        priority = url_el.findtext("sm:priority", default="0", namespaces=namespace)

        try:
            last_modified = datetime.fromisoformat(lastmod.replace("Z", "+00:00"))
        except ValueError:
            last_modified = datetime.now()

        parts = urlparse(loc).path.strip("/").split("/")
        section = parts[1] if len(parts) > 1 else "unknown"
        date = parts[2] if len(parts) > 2 else ""
        slug = parts[3] if len(parts) > 3 else ""

        entries.append(
            SitemapEntry(
                url=loc,
                last_modified=last_modified,
                change_freq=changefreq,
                priority=float(priority),
                section=section,
                date=date,
                slug=slug,
            )
        )

    return entries


def get_bloomberg_urls() -> list[str]:
    """Pull {_COUNTRY}-related URLs from the current month's Bloomberg sitemap."""
    now = datetime.now()
    xml_url = f"https://www.bloomberg.com/sitemaps/news/{now.year}-{now.month}.xml"
    tprint(f"[Bloomberg] Fetching sitemap: {xml_url}")
    try:
        entries = parse_bloomberg_sitemap(xml_url)
        search_country = _COUNTRY.lower().replace(" ", "-")
        urls = [e.url for e in entries if search_country in e.url.lower()]
        tprint(f"[Bloomberg] Found {len(urls)} {_COUNTRY} articles in sitemap")
        return urls
    except Exception as e:
        tprint(f"[Bloomberg] Sitemap error: {e}")
        return []


# ─────────────────────────────────────────────
#  STEP 1B — Gemini web search for other sources
# ─────────────────────────────────────────────
def get_urls_from_gemini(sources: list[str], min_per_source: int = 10) -> list[str]:
    """
    Ask Gemini to search for the latest {_COUNTRY} political news URLs
    from the specified sources, returning at least min_per_source per source.
    """
    sources_list = ", ".join(sources)

    prompt = f"""
Today's date is {_MONTH_YEAR}. The country you need to search for is {_COUNTRY}.

Search the web and return a list of URLs to the most recent and politically relevant
{_COUNTRY} news articles published in the last 30 days from these sources:
{SEARCH_SOURCES}

Requirements:
- Return AT LEAST {min_per_source} URLs per source
- Articles must be specifically about {_COUNTRY}'s POLITICS:
    * Internal politics (elections, constitutional reform, government policy,
      nuclear energy, economic policy, party dynamics)
    * External politics (the shifts in its relationships, strategic alliances, and actions on the global stage, often driven by both domestic interests and international dynamics, trade etc)
    * {_COUNTRY}'s relationship with the UAE — energy, defence, diplomacy, trade etc. 
- Only return real, working article URLs (not homepage or section-index URLs)
- Include the most impactful and widely-covered stories

Return ONLY a JSON object with this exact structure:
{{
  "urls": ["https://...", "https://...", ...]
}}
No explanation, no markdown, just the JSON.
"""

    tprint(f"[Gemini Search] Querying for URLs from: {sources_list}")
    try:
        response = _client.models.generate_content(
            model=GEMINI_MODEL,
            contents=prompt,
            config=_search_config,
        )
        raw = response.text.strip().lstrip("```json").rstrip("```").strip()
        data = json.loads(raw)
        urls = data.get("urls", [])
        tprint(f"[Gemini Search] Returned {len(urls)} URLs")
        return urls
    except Exception as e:
        tprint(f"[Gemini Search] Error: {e}")
        return []


# ─────────────────────────────────────────────
#  STEP 1 — Collect all URLs
# ─────────────────────────────────────────────
def collect_urls(country: str) -> list[str]:
    """Combine Bloomberg sitemap URLs + Gemini-searched URLs from other sources."""
    bloomberg_urls = get_bloomberg_urls()
    other_urls = get_urls_from_gemini(sources=SEARCH_SOURCES, min_per_source=10)
    tprint(f"Bloomberg URLs: {bloomberg_urls}")
    tprint(f"Other URLs: {other_urls}")
    all_urls = list(
        dict.fromkeys(bloomberg_urls + other_urls)
    )  # deduplicate, preserve order
    tprint(
        f"\n[Collect] Total unique URLs: {len(all_urls)} "
        f"({len(bloomberg_urls)} Bloomberg + {len(other_urls)} other)\n"
    )
    return all_urls


# ─────────────────────────────────────────────
#  Archive URL builder — always use archive.today
# ─────────────────────────────────────────────
def get_archive_url(url: str) -> str | None:
    """Always route through archive.today. Skip blocked domains."""
    if any(domain in url for domain in BLOCKED_DOMAINS):
        return None
    match = re.search(r"https?://[^\s]+", url)
    if not match:
        return None
    return f"https://archive.today/newest/{match.group()}"


# ─────────────────────────────────────────────
#  HTTP helpers
# ─────────────────────────────────────────────
def get_headers() -> dict:
    return {
        "User-Agent": random.choice(USER_AGENTS),
        "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,image/webp,*/*;q=0.8",
        "Accept-Language": "en-US,en;q=0.9",
        "Accept-Encoding": "gzip, deflate, br",
        "Connection": "keep-alive",
        "Upgrade-Insecure-Requests": "1",
        "Cache-Control": "max-age=0",
        "Referer": "https://www.google.com/",
    }


def fetch_with_retry(
    url: str, max_retries: int = MAX_RETRIES, backoff_base: float = 2.0
) -> requests.Response:
    """Fetch with exponential backoff. Coordinates via shared rate limiter + global backoff."""
    session = requests.Session()

    for attempt in range(1, max_retries + 1):
        _global_backoff.wait_if_needed()
        _rate_limiter.wait()

        try:
            tprint(f"  [fetch] Attempt {attempt}/{max_retries}: {url}")
            response = session.get(
                url, headers=get_headers(), timeout=30, allow_redirects=True
            )

            if response.status_code == 429:
                retry_after = int(
                    response.headers.get("Retry-After", backoff_base**attempt)
                )
                wait = retry_after + random.uniform(2, 5)
                _global_backoff.trigger(wait)
                continue

            response.raise_for_status()
            return response

        except requests.exceptions.HTTPError as e:
            if attempt == max_retries:
                raise
            wait = backoff_base**attempt + random.uniform(0, 2)
            tprint(f"  HTTP error ({e}) — retrying in {wait:.1f}s...")
            time.sleep(wait)

        except requests.exceptions.RequestException as e:
            if attempt == max_retries:
                raise
            wait = backoff_base**attempt
            tprint(f"  Request error ({e}) — retrying in {wait:.1f}s...")
            time.sleep(wait)

    raise RuntimeError(f"Failed to fetch {url} after {max_retries} attempts.")


# ─────────────────────────────────────────────
#  BeautifulSoup scraper (also a Gemini tool)
# ─────────────────────────────────────────────
def scrape_with_beautifulsoup(url: str) -> str:
    """Fetch a page and return clean article text."""
    response = fetch_with_retry(url)
    soup = BeautifulSoup(response.text, "html.parser")

    for tag in soup(
        [
            "script",
            "style",
            "noscript",
            "nav",
            "footer",
            "header",
            "aside",
            "form",
            "iframe",
        ]
    ):
        tag.decompose()

    body_text = ""
    candidates = [
        soup.find("article"),
        soup.find(id="article-content"),
        soup.find(class_="article-body"),
        soup.find(class_="body-content"),
        soup.find(class_="post-content"),
        soup.find("main"),
    ]

    for candidate in candidates:
        if candidate:
            paragraphs = candidate.find_all("p")
            text = "\n\n".join(
                p.get_text(strip=True) for p in paragraphs if p.get_text(strip=True)
            )
            if len(text) > 200:
                body_text = text
                break

    if not body_text:
        paragraphs = soup.find_all("p")
        body_text = "\n\n".join(
            p.get_text(strip=True) for p in paragraphs if p.get_text(strip=True)
        )

    return body_text


# ─────────────────────────────────────────────
#  Tool dispatcher
# ─────────────────────────────────────────────
_TOOL_REGISTRY: dict[str, callable] = {
    "scrape_with_beautifulsoup": lambda args: scrape_with_beautifulsoup(args["url"]),
}


def _dispatch_tool_call(fn_name: str, fn_args: dict) -> str:
    handler = _TOOL_REGISTRY.get(fn_name)
    if not handler:
        return f"Error: unknown tool '{fn_name}'"
    try:
        result = handler(fn_args)
        return result if isinstance(result, str) else json.dumps(result)
    except Exception as e:
        return f"Error in {fn_name}: {e}"


# ─────────────────────────────────────────────
#  STEP 2 — Single-URL scraper (runs in each thread)
#
#  Two-pass strategy (Gemini forbids mixing url_context + FunctionDeclarations):
#    Pass 1 — url_context only config: Gemini reads the page directly.
#    Pass 2 — BS4 function tool config: called only if Pass 1 returns no content.
#              Gemini drives the tool-use loop, we execute the function locally.
# ─────────────────────────────────────────────
def _run_tool_loop(response, messages: list) -> str:
    """
    Execute Gemini's function-call loop using _scrape_bs4_config.
    Keeps sending tool results back until Gemini returns plain text.
    Returns the final text content, or empty string if nothing returned.
    """
    while True:
        fn_calls = [
            part
            for candidate in response.candidates
            for part in candidate.content.parts
            if hasattr(part, "function_call") and part.function_call
        ]

        if not fn_calls:
            break

        # Append Gemini's model turn (the function_call request)
        messages.append(
            {
                "role": "model",
                "parts": [
                    {
                        "function_call": {
                            "name": fc.function_call.name,
                            "args": dict(fc.function_call.args),
                        }
                    }
                    for fc in fn_calls
                ],
            }
        )

        # Execute each tool locally and collect results
        tool_results = []
        for fc in fn_calls:
            fn_name = fc.function_call.name
            fn_args = dict(fc.function_call.args)
            tprint(f"  [tool] Gemini → '{fn_name}({fn_args})'")
            result_text = _dispatch_tool_call(fn_name, fn_args)
            tool_results.append(
                {
                    "function_response": {
                        "name": fn_name,
                        "response": {"content": result_text},
                    }
                }
            )

        # Send tool results back and get the next response
        messages.append({"role": "user", "parts": tool_results})
        response = _client.models.generate_content(
            model=GEMINI_MODEL,
            contents=messages,
            config=_scrape_bs4_config,  # stay on BS4 config for follow-up turns
        )

    return response.text or ""


def scrape_page(url: str) -> dict:
    """
    Scrape one URL. Never raises — errors returned in dict.

    Pass 1: Ask Gemini to read the archive.today URL using url_context only.
    Pass 2: If Pass 1 yields no content, switch to a separate request with
            the BS4 FunctionDeclaration tool so Gemini can call scrape_with_beautifulsoup.
    """
    archive_url = get_archive_url(url)
    if not archive_url:
        return {
            "url": url,
            "archive_url": None,
            "content": None,
            "error": "Blocked or unrecognised URL",
        }

    base_prompt = f"""Extract the main article text from this URL.
    Use url_context to retrieve content from the page if it is accessible.
    Return JSON with keys: "title" (string) and "body" (string).
    Do not fabricate content — if the page is inaccessible return empty strings.

URL: {archive_url}"""

    for attempt in range(1, MAX_RETRIES + 1):
        try:
            # ── Pass 1: url_context only ──────────────────────────────────
            tprint(f"  [Pass 1 / url_context] {archive_url}")
            response = _client.models.generate_content(
                model=GEMINI_MODEL,
                contents=base_prompt,
                config=_scrape_url_context_config,
            )
            content = response.text or ""

            # Check if Pass 1 actually got article content
            has_content = False
            if content:
                try:
                    parsed = json.loads(content)
                    has_content = bool(parsed.get("body", "").strip())
                except Exception:
                    has_content = len(content.strip()) > 100
            time.sleep(random.uniform(1, 5))
            # ── Pass 2: BS4 function tool (only if Pass 1 empty) ─────────
            if not has_content:
                tprint(
                    f"  [Pass 2 / BS4 tool] url_context returned nothing — switching to BS4 tool"
                )
                bs4_prompt = f"""
                Call the scrape_with_beautifulsoup tool or use BeautifulSoup with the URL below to fetch and extract the article text.
                Return JSON with keys: "title" (string) and "body" (string).
                Do not fabricate content.

                URL: {archive_url}"""

                response = _client.models.generate_content(
                    model=GEMINI_MODEL,
                    contents=bs4_prompt,
                    config=_scrape_bs4_config,
                )
                content = response.text or ""

            if not content:
                tprint(f"  ✘ Failed: [{url}] no content after both passes")
                return {
                    "url": url,
                    "archive_url": archive_url,
                    "content": None,
                    "error": "No content found",
                }

            tprint(f"  ✔ Done: {url}")
            return {
                "url": url,
                "archive_url": archive_url,
                "content": content,
                "error": None,
            }

        except Exception as e:
            tprint(f"  ✘ [{url}] attempt {attempt}/{MAX_RETRIES}: {e}")
            if attempt < MAX_RETRIES:
                time.sleep(RETRY_DELAY * attempt)
            else:
                return {
                    "url": url,
                    "archive_url": archive_url,
                    "content": None,
                    "error": str(e),
                }


# ─────────────────────────────────────────────
#  Multithreaded scraper
# ─────────────────────────────────────────────
def scrape_all(urls: list[str], max_workers: int = MAX_WORKERS) -> list[dict]:
    results = []
    total = len(urls)
    tprint(f"\n[Scrape] {total} URLs across {max_workers} threads\n")

    with ThreadPoolExecutor(max_workers=max_workers) as executor:
        future_to_url = {executor.submit(scrape_page, url): url for url in urls}
        for i, future in enumerate[Future[dict[Any, Any]]](
            as_completed(future_to_url), start=1
        ):
            url = future_to_url[future]
            result = future.result()
            tprint(f"  [{i}/{total}] {'✔' if not result['error'] else '✘'} {url}")
            results.append(result)

    return results


# ─────────────────────────────────────────────
#  STEP 3 — Report generator
# ─────────────────────────────────────────────
def report_generator(country: str, scraped_results: list[dict]) -> dict:
    """
    Synthesise all scraped articles into a structured analytical report
    matching the format of the {country} Country Report template:
      - Overview (bold analytical summary paragraph)
      - UAE relevance paragraph
      - 1. Internal Developments  (3 sub-sections, each with heading + body)
      - 2. External Developments  (3 sub-sections, do not include any UAE relevance paragraphs)
      - 3. Of Relevance to the UAE (1–2 paragraphs, summarise all the UAE relevance news, bold key sentences)

    Returns a dict with the report structure ready for docx generation.
    # Style Guide(MUST FOLLOW):

    - Use British English consistently. Example:Write "Emphasise" instead of "emphasize". Write "labour" instead of "labor"
    - Abbreviations and Acronyms: Write out abbreviations fully.Use "US" and "UN" (not "U.S." or "U.N.")
    - Currency amounts: Format as "EUR XXm" or "USD XXbn". No currency symbols.
    - Use "%" symbol instead of writing "percent"
    - Quotations: Longer than half a line: in italics. Short quotes: no special formatting
    - Aim for a punchy, journalistic approach.
    - Prioritise clarity and conciseness.
    - Avoid overly detailed information. Use generic descriptors (e.g., "CEO of X company").Focus on relative variations over absolute numbers
    - Date format: Month DDth (e.g., December 12th)
    """

    # Build a condensed corpus from successful scrapes
    articles_text = ""
    for r in scraped_results:
        if r.get("content") and not r.get("error"):
            try:
                parsed = json.loads(r["content"])
                title = parsed.get("title", "")
                body = parsed.get("body", "")[:1500]  # cap per article
                articles_text += (
                    f"\n\n--- SOURCE: {r['url']} ---\nTITLE: {title}\n{body}"
                )
            except Exception:
                articles_text += (
                    f"\n\n--- SOURCE: {r['url']} ---\n{str(r['content'])[:1500]}"
                )

    prompt = f"""
You are a senior geopolitical analyst writing a confidential country brief for {country} on {_MONTH_YEAR}.

Below are raw news articles about {country} politics scraped from {SEARCH_SOURCES}.

Your task: synthesise these into a structured analytical report in the EXACT format below.
Write in a sharp, analytical, intelligence-brief style — not a plain summary.
Highlight strategic implications, signal what is driving key actors, and note what matters for the UAE.
Bold the most strategically significant sentences.

# MUST FOLLOW THESE RULES:
- Make sure to use British English consistently. Example:Write "Emphasise" instead of "emphasize". Write "labour" instead of "labor".
- Abbreviations and Acronyms: Write out abbreviations fully.Use "US" and "UN" (not "U.S." or "U.N.")
- Currency amounts: Format as "EUR XXm" or "USD XXbn". No currency symbols.
- Use "%" symbol instead of writing "percent"
- Quotations: Longer than half a line: in italics. Short quotes: no special formatting
- Aim for a punchy, journalistic approach.
- Prioritise clarity and conciseness.
- Avoid overly detailed information. Use generic descriptors (e.g., "CEO of X company").Focus on relative variations over absolute numbers
- Date format: Month DDth (e.g., December 12th)

REQUIRED OUTPUT FORMAT (return as JSON):
{{
  "title": "Country Report — {_COUNTRY}\\t\\t\\t\\t    {_MONTH_YEAR}",
  "overview": "Single analytical paragraph summarising the most important strategic developments. Bold the 1-2 most critical sentences using **markdown bold**.",
  "uae_overview": "Single bold paragraph on what the top {country} story means for the UAE specifically.",
  "internal_developments": [
    {{
      "heading": "Short sub-heading (max 12 words)",
      "body": "2-4 sentence analytical paragraph. **Bold** the most important sentence."
    }},
    ... (exactly 3 items)
  ],
  "external_developments": [
    {{
      "heading": "Short sub-heading (max 12 words)",
      "body": "2-4 sentence analytical paragraph. **Bold** the most important sentence."
    }},
    ... (exactly 3 items)
  ],
  "uae_relevance": [
    {{
      "heading": "Short sub-heading (max 12 words)",
      "body": "2-4 sentence analytical paragraph. **Bold** the most important sentence."
    }},
    ... (exactly 3 items)
  ],
}}

ARTICLES:
{articles_text}
"""

    tprint("\n[Report] Generating analytical report via Gemini...")
    try:
        response = _client.models.generate_content(
            model=GEMINI_MODEL,
            contents=prompt,
            config=_report_config,
        )
        raw = response.text.strip().lstrip("```json").rstrip("```").strip()
        report = json.loads(raw)
        tprint("[Report] Report structure generated successfully")
        return report
    except Exception as e:
        tprint(f"[Report] Error generating report: {e}")
        raise


# ─────────────────────────────────────────────
#  STEP 4 — Save report as .docx
# ─────────────────────────────────────────────
def get_report_path() -> Path:
    """Return a versioned report path, appending _v2, _v3 etc. if today's file exists."""
    REPORTS_DIR.mkdir(parents=True, exist_ok=True)
    date_str = datetime.now().strftime("%Y-%m-%d")
    base = REPORTS_DIR / f"COUNTRY REPORT - {_COUNTRY} - March 2026"
    path = base.with_suffix(".docx")
    if not path.exists():
        return path
    v = 2
    while True:
        path = base.parent / f"{base.name}_v{v}.docx"
        if not path.exists():
            return path
        v += 1


def _parse_bold_runs(text: str) -> list[dict]:
    """
    Convert a string with **bold** markers into a list of run dicts:
    [{"text": "...", "bold": False}, {"text": "...", "bold": True}, ...]
    """
    runs = []
    pattern = re.compile(r"\*\*(.+?)\*\*", re.DOTALL)
    last = 0
    for m in pattern.finditer(text):
        if m.start() > last:
            runs.append({"text": text[last : m.start()], "bold": False})
        runs.append({"text": m.group(1), "bold": True})
        last = m.end()
    if last < len(text):
        runs.append({"text": text[last:], "bold": False})
    return runs or [{"text": text, "bold": False}]


# ─────────────────────────────────────────────
#  STEP 4 — Save report as .docx (python-docx)
# ─────────────────────────────────────────────
def _add_runs(para, run_list: list, bold_override: bool = False, size_pt: float = 13):
    for r in run_list:
        is_bold = bold_override or r.get("bold", False)
        run = para.add_run(r["text"])
        run.bold = is_bold
        run.font.name = _FONT_MEDIUM
        run.font.size = Pt(size_pt)


def _add_section_heading(doc, num: str, label: str):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(16)
    para.paragraph_format.space_after = Pt(8)
    run = para.add_run(f"{num}.  {label}")
    run.font.name = "Lora"
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)


def _add_sub_heading(doc, num_str: str, heading_text: str):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(10)
    para.paragraph_format.space_after = Pt(5)
    run = para.add_run(f"{num_str}  {heading_text}")
    run.font.name = "Lora"
    run.font.size = Pt(13)
    run.underline = True


def _add_body_para(doc, run_list: list, bold_override: bool = False):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    para.paragraph_format.space_after = Pt(8)
    _add_runs(para, run_list, bold_override=bold_override)


def save_report_docx(report: dict) -> Path:
    """Generate a .docx file matching the {_COUNTRY} Country Report template.
    Pure python-docx — no Node.js or subprocess required.
    """
    out_path = get_report_path()

    # Defensive defaults — Gemini occasionally returns unexpected shapes
    for key in ("internal_developments", "external_developments", "uae_relevance"):
        for item in report.get(key, []):
            item.setdefault("heading", "")
            item.setdefault("body", "")

    # uae_relevance = report.get("uae_relevance", "")
    # if isinstance(uae_relevance, list):
    #     uae_relevance = "\n\n".join(uae_relevance)

    doc = DocxDocument()

    # Page setup: US Letter, 1-inch margins
    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    sec.left_margin = sec.right_margin = Inches(1)
    sec.top_margin = sec.bottom_margin = Inches(1)

    # Title
    title_para = doc.add_paragraph()
    title_para.paragraph_format.space_after = Pt(16)
    run = title_para.add_run(
        report.get(
            "title", f"Country Report — {_COUNTRY}\t\t\t\t" + " " * 4 + _MONTH_YEAR
        )
    )
    run.bold = True
    run.font.name = "Lora"
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)

    # Overview heading
    ov_head = doc.add_paragraph()
    ov_head.paragraph_format.space_after = Pt(6)
    run = ov_head.add_run("Overview")
    run.font.name = "Lora"
    run.font.size = Pt(14)

    # Overview body + UAE overview (bold)
    _add_body_para(doc, _parse_bold_runs(report.get("overview", "")))
    # _add_body_para(
    #     doc, _parse_bold_runs(report.get("uae_overview", "")), bold_override=True
    # )

    # 1. Internal Developments
    _add_section_heading(doc, "1", "Internal Developments")
    for i, item in enumerate(report.get("internal_developments", []), 1):
        _add_sub_heading(doc, f"1.{i}.", item["heading"])
        _add_body_para(doc, _parse_bold_runs(item["body"]))

    # 2. External Developments
    _add_section_heading(doc, "2", "External Developments")
    for i, item in enumerate(report.get("external_developments", []), 1):
        _add_sub_heading(doc, f"2.{i}.", item["heading"])
        _add_body_para(doc, _parse_bold_runs(item["body"]))

    # 3. Of Relevance to the UAE
    _add_section_heading(doc, "3", "Of Relevance to the UAE")
    for i, item in enumerate(report.get("external_developments", []), 1):
        _add_sub_heading(doc, f"3.{i}.", item["heading"])
        _add_body_para(doc, _parse_bold_runs(item["body"]))

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    tprint(f"[Docx] Saved → {out_path}")
    return out_path


# ─────────────────────────────────────────────
#  Main pipeline
# ─────────────────────────────────────────────
def main():
    try:
        with open(f"scraped_results_{_COUNTRY}.json", "r") as f:
            scraped_results = json.load(f)
    except FileNotFoundError:
        scraped_results = None
    if not scraped_results:
        # Step 1 — collect URLs
        tprint("=" * 60)
        tprint("STEP 1: Collecting URLs")
        tprint("=" * 60)
        urls = collect_urls(country=_COUNTRY)

        # Step 2 — scrape all URLs concurrently
        tprint("=" * 60)
        tprint("STEP 2: Scraping articles")
        tprint("=" * 60)
        scraped_results = scrape_all(urls, max_workers=MAX_WORKERS)

        ok = sum(1 for r in scraped_results if not r["error"])
        fail = len(scraped_results) - ok
        tprint(
            f"\n[Scrape] {ok} succeeded, {fail} failed out of {len(scraped_results)} total"
        )

        # Save raw scrape results for debugging
        with open(f"scraped_results_{_COUNTRY}.json", "w") as f:
            json.dump(scraped_results, f, indent=2, default=str)
        tprint("[Scrape] Raw results saved → scraped_results.json")

    # Step 3 — generate report
    tprint("=" * 60)
    tprint("STEP 3: Generating report")
    tprint("=" * 60)
    report = report_generator(country=_COUNTRY, scraped_results=scraped_results)

    with open("report_structure.json", "w") as f:
        json.dump(report, f, indent=2, default=str)
    tprint("[Report] Report structure saved → report_structure.json")

    # Step 4 — save docx
    tprint("=" * 60)
    tprint("STEP 4: Saving .docx")
    tprint("=" * 60)
    docx_path = save_report_docx(report)
    tprint(f"\n✔ Report saved → {docx_path}")

    return docx_path


if __name__ == "__main__":
    main()
